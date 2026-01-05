import pandas as pd
import numpy as np
import statsmodels.api as sm
from docx import Document
from docx.shared import Inches

# CONFIG (change paths if needed)
EPI_CSV = "epi_quarterly_true.csv"
FIN_CSV = "mbs_quarterly_metrics.csv"

OUT_MERGED = "mbs_merged_epi_financial.csv"
OUT_RESULTS = "regression_results_long.csv"
OUT_TABLE_DOCX = "apa_regression_table.docx"

DEP_VAR = "revpar"   # dependent variable in financial file

# HELPERS
def to_period_q(s: str) -> pd.Period:
    """Convert '2015Q1' -> Period('2015Q1', 'Q')"""
    return pd.Period(str(s), freq="Q")

def add_quarter_dummies(df: pd.DataFrame) -> pd.DataFrame:
    """
    Add seasonal control, avoiding dummy trap.

    IMPORTANT: Your lagging removes Q1 observations (no q-1 because Q4 missing),
    so the regression sample typically contains only Q2 & Q3.
    In that case, if you include both Q2 and Q3 dummies + intercept, you get
    perfect collinearity because (Q2 + Q3) == 1 for every row.

    Solution: include ONLY one dummy, e.g. Q3 (Q2 is baseline).
    """
    qnum = df["quarter_p"].dt.quarter
    df["Q3"] = (qnum == 3).astype(int)
    return df

def run_ols_hc3(y, X):
    """OLS with HC3 robust SE."""
    X = sm.add_constant(X, has_constant="add")
    model = sm.OLS(y, X, missing="drop")
    res = model.fit(cov_type="HC3")
    return res

def format_p(p):
    if p < 0.001:
        return "< .001"
    return f"{p:.3f}"

def stars(p):
    if p < 0.001:
        return "***"
    if p < 0.01:
        return "**"
    if p < 0.05:
        return "*"
    return ""

def apa_table_to_docx(res_list, model_names, filename):
    """
    Create a simple APA-style regression table in Word.
    Shows b (SE) with significance stars.
    """
    # Collect all coefficient names across models
    coef_names = set()
    for r in res_list:
        coef_names |= set(r.params.index)
    # Order: constant first, then EPI_lag1, then controls
    preferred_order = ["const", "EPI_lag1", "Q3", "occupancy", "trevpar"]
    remaining = [c for c in coef_names if c not in preferred_order]
    coef_order = [c for c in preferred_order if c in coef_names] + sorted(remaining)

    doc = Document()
    doc.add_heading("Regression Results (APA style)", level=1)

    # Table: rows = coefficients + model stats, cols = models+1
    table = doc.add_table(rows=1, cols=1 + len(res_list))
    hdr = table.rows[0].cells
    hdr[0].text = "Predictor"
    for j, name in enumerate(model_names, start=1):
        hdr[j].text = name

    # Coeff rows
    for c in coef_order:
        row = table.add_row().cells
        row[0].text = "Intercept" if c == "const" else c
        for j, r in enumerate(res_list, start=1):
            if c in r.params.index:
                b = r.params[c]
                se = r.bse[c]
                p = r.pvalues[c]
                row[j].text = f"{b:.3f} ({se:.3f}){stars(p)}"
            else:
                row[j].text = ""

    # Add model fit stats
    def add_stat(label, values):
        row = table.add_row().cells
        row[0].text = label
        for j, v in enumerate(values, start=1):
            row[j].text = v

    add_stat("N", [str(int(r.nobs)) for r in res_list])
    add_stat("R²", [f"{r.rsquared:.3f}" for r in res_list])
    add_stat("Adj. R²", [f"{r.rsquared_adj:.3f}" for r in res_list])
    add_stat("F (p)", [f"{r.fvalue:.2f} ({format_p(r.f_pvalue)})" if r.fvalue is not None else "" for r in res_list])

    doc.add_paragraph("\nNote. Entries are unstandardized coefficients with HC3 robust standard errors in parentheses.")
    doc.add_paragraph("Significance: * p < .05, ** p < .01, *** p < .001.")
    doc.save(filename)

# LOAD + MERGE
epi = pd.read_csv(EPI_CSV)
fin = pd.read_csv(FIN_CSV)

# Standardize quarter types
epi["quarter_p"] = epi["quarter"].apply(to_period_q)
fin["quarter_p"] = fin["quarter"].apply(to_period_q)

# Keep only necessary columns
epi = epi[["quarter", "quarter_p", "EPI", "n_reviews"]].copy()

# Financial file may have occupancy/trevpar missing -> keep if present
keep_fin_cols = ["quarter", "quarter_p", "revpar"]
for c in ["occupancy", "trevpar"]:
    if c in fin.columns:
        keep_fin_cols.append(c)

fin = fin[keep_fin_cols].copy()

# Merge on quarter
df = pd.merge(fin, epi, on=["quarter", "quarter_p"], how="inner")

# LAGGED EPI (true q-1)
# Create EPI_{q-1} by shifting period by +1 and merging back:
# For each quarter q, we want EPI from (q-1).
epi_lag = epi.copy()
epi_lag["quarter_p"] = epi_lag["quarter_p"] + 1  # move forward one quarter
epi_lag = epi_lag.rename(columns={"EPI": "EPI_lag1", "n_reviews": "n_reviews_lag1"})

df = pd.merge(df, epi_lag[["quarter_p", "EPI_lag1", "n_reviews_lag1"]], on="quarter_p", how="left")

# Add seasonal dummy (Q3 only to avoid dummy trap)
df = add_quarter_dummies(df)

# Save merged dataset
df_out = df.sort_values("quarter_p").copy()
df_out.to_csv(OUT_MERGED, index=False)
print(f"Saved merged dataset: {OUT_MERGED}")

# REGRESSIONS
# Drop rows without true lag
reg = df_out.dropna(subset=["EPI_lag1", DEP_VAR]).copy()

# Baseline model: RevPAR ~ EPI_lag1 + Q3
X_base = reg[["EPI_lag1", "Q3"]].copy()
y = reg[DEP_VAR].astype(float)
res_base = run_ols_hc3(y, X_base)

# Extended model: add occupancy + trevpar if available
X_ext_cols = ["EPI_lag1", "Q3"]
if "occupancy" in reg.columns:
    X_ext_cols.append("occupancy")
if "trevpar" in reg.columns:
    X_ext_cols.append("trevpar")

X_ext = reg[X_ext_cols].copy()
res_ext = run_ols_hc3(y, X_ext)

print("\nBaseline model summary (HC3):")
print(res_base.summary())
print("\nExtended model summary (HC3):")
print(res_ext.summary())

# EXPORT RESULTS (CSV + APA DOCX)
def results_to_long(res, model_name):
    rows = []
    for name in res.params.index:
        rows.append({
            "model": model_name,
            "term": name,
            "b": res.params[name],
            "se_hc3": res.bse[name],
            "t": res.tvalues[name],
            "p": res.pvalues[name],
        })
    # model stats
    rows.append({"model": model_name, "term": "N", "b": res.nobs, "se_hc3": np.nan, "t": np.nan, "p": np.nan})
    rows.append({"model": model_name, "term": "R2", "b": res.rsquared, "se_hc3": np.nan, "t": np.nan, "p": np.nan})
    rows.append({"model": model_name, "term": "Adj_R2", "b": res.rsquared_adj, "se_hc3": np.nan, "t": np.nan, "p": np.nan})
    return rows

long = pd.DataFrame(results_to_long(res_base, "Baseline") + results_to_long(res_ext, "Extended"))
long.to_csv(OUT_RESULTS, index=False)
print(f"Saved regression results (long CSV): {OUT_RESULTS}")

apa_table_to_docx([res_base, res_ext], ["Model 1 (Baseline)", "Model 2 (Extended)"], OUT_TABLE_DOCX)
print(f"Saved APA-style regression table: {OUT_TABLE_DOCX}")

# SANITY PRINTS
print("\nRegression sample quarters used:")
print(reg[["quarter", "EPI_lag1", DEP_VAR]].head(10))
print("\nN baseline:", int(res_base.nobs), " | N extended:", int(res_ext.nobs))
print("Note: Q1 observations drop out because Q4 is not available to create a true q−1 lag.")
