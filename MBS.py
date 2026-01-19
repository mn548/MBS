
# Purpose:
#   This script merges EPI (Economic Policy Uncertainty) data
#   with financial / MBS data, engineers quarterly metrics,
#   and prepares a modeling dataset for downstream regression
#   analysis.
#
#   It also runs baseline and extended regression models and
#   exports results to CSV and an APA-style Word table.
#
# Expected inputs (CSV files in repo root):
#   - epi_quarterly_true.csv
#   - epi_topics.csv
#   - epi_aspect_quarterly.csv
#   - Financial data (loaded inside Financial data.py or here)
#
# Outputs (written to repo root):
#   - mbs_quarterly_metrics.csv
#   - mbs_merged_epi_financial.csv
#   - regression_results_long.csv
#   - apa_regression_table.docx
#
# How to run:
#   python MBS.py

# LOAD ALL NECESSARY LIBRARIES
import pandas as pd
import numpy as np

# NumPy backward-compatibility shim
_ALIAS_MAP = {
    "complex_": np.complex128,
    "float_": np.float64,
    "string_": np.bytes_,
    "unicode_": np.str_,
}

for name, dtype in _ALIAS_MAP.items():
    if not hasattr(np, name):
        setattr(np, name, dtype)

for name, dtype in _ALIAS_MAP.items():
    if name not in np.sctypeDict:
        np.sctypeDict[name] = dtype

from pathlib import Path
import re
import statsmodels.api as sm
from docx import Document

# GLOBAL PATHS AND CONSTANTS
BASE_DIR = Path(__file__).resolve().parent

# Input file paths
EPI_QUARTERLY = BASE_DIR / "epi_quarterly_true.csv"
EPI_TOPICS = BASE_DIR / "epi_topics.csv"
EPI_ASPECTS = BASE_DIR / "epi_aspect_quarterly.csv"

# Output file paths
OUT_MBS_QTR = BASE_DIR / "mbs_quarterly_metrics.csv"
OUT_MERGED = BASE_DIR / "mbs_merged_epi_financial.csv"
OUT_RESULTS = BASE_DIR / "regression_results_long.csv"
OUT_TABLE_DOCX = BASE_DIR / "apa_regression_table.docx"

# Dependent variable name used in regressions
DEP_VAR = "mbs_return"

# HELPER DICTIONARIES
# Month abbreviation normalization
MONTH_ABBR = {
    "Jan": "January", "Feb": "February", "Mar": "March",
    "Apr": "April", "May": "May", "Jun": "June",
    "Jul": "July", "Aug": "August", "Sep": "September",
    "Oct": "October", "Nov": "November", "Dec": "December",
}

# Month-to-quarter mapping
MONTH_TO_Q = {
    "January": 1, "February": 1, "March": 1,
    "April": 2, "May": 2, "June": 2,
    "July": 3, "August": 3, "September": 3,
    "October": 4, "November": 4, "December": 4,
}

# TEXT AND DATE PARSING UTILS
def flatten_text(df: pd.DataFrame) -> str:
    """
    Convert an entire DataFrame into a single lowercase text string.

    Used for extracting quarter information from messy or
    multi-cell textual financial disclosures.
    """
    return " ".join(df.astype(str).values.flatten()).lower()


def infer_quarter_from_text(df: pd.DataFrame) -> str | None:
    """
    Attempt to infer a fiscal quarter string (e.g., '2022Q2')
    from unstructured financial text.

    Strategy:
      1) Look for patterns like 'Three months ended Jun 30, 2022'
      2) Fallback to 'Period End: Jun 30, 2022'

    Returns:
      - Quarter string like '2022Q2'
      - None if no valid quarter could be inferred
    """
    text = flatten_text(df)

    # Pattern: 'three months ended Jun 30, 2022'
    m = re.search(r"three months ended\s+([A-Za-z]+)\s+(\d{1,2})(?:,)?\s+(\d{4})", text, flags=re.IGNORECASE)
    if m:
        mon, _day, year = m.groups()
        mon = mon.strip().capitalize()
        mon = MONTH_ABBR.get(mon[:3].capitalize(), mon)
        q = MONTH_TO_Q.get(mon)
        if q:
            return f"{year}Q{q}"

    # Fallback pattern: 'period end: Jun 30, 2022'
    m = re.search(r"period end\s*[:]?\s*([A-Za-z]+)\s+(\d{1,2})(?:,)?\s+(\d{4})", text, flags=re.IGNORECASE)
    if m:
        mon, _day, year = m.groups()
        mon = mon.strip().capitalize()
        mon = MONTH_ABBR.get(mon[:3].capitalize(), mon)
        q = MONTH_TO_Q.get(mon)
        if q:
            return f"{year}Q{q}"

    return None

# DATA LOADING FUNCTIONS
def load_epi_data() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    """
    Load all EPI-related datasets.

    Returns:
      - Quarterly EPI series
      - Topic-level EPI data
      - Aspect-level quarterly EPI data
    """
    epi_q = pd.read_csv(EPI_QUARTERLY)
    epi_topics = pd.read_csv(EPI_TOPICS)
    epi_aspects = pd.read_csv(EPI_ASPECTS)

    # Ensure consistent column naming
    for df in (epi_q, epi_topics, epi_aspects):
        df.columns = df.columns.str.strip()

    return epi_q, epi_topics, epi_aspects


def load_financial_data() -> pd.DataFrame:
    """
    Placeholder loader for financial / MBS data.

    In the current workflow, this may be loaded in
    'Financial data.py' and saved to disk for reuse.
    Modify this function if you want MBS.py to be
    fully self-contained.
    """
    # Example: replace with real loading logic
    df = pd.read_csv(BASE_DIR / "mbs_financial.csv")
    return df

# FEATURE ENGINEERING
def build_quarterly_metrics(epi_q: pd.DataFrame, fin: pd.DataFrame) -> pd.DataFrame:
    """
    Merge EPI and financial data and compute quarterly metrics.

    Steps:
      - Align on quarter
      - Create lagged EPI variables
      - Compute returns or growth rates for MBS

    Returns:
      - DataFrame with engineered quarterly features
    """
    # Merge on quarter key
    df = pd.merge(fin, epi_q, on="quarter", how="inner")

    # Sort by time for lag creation
    df = df.sort_values("quarter")

    # Lagged EPI (1 quarter)
    df["EPI_lag1"] = df["EPI"].shift(1)

    # Example: MBS quarterly return
    df["mbs_return"] = df["mbs_price"].pct_change()

    # Drop rows with missing values caused by lags
    df = df.dropna().reset_index(drop=True)

    return df

# REGRESSION UTILITIES
def run_regression(df: pd.DataFrame, x_vars: list[str], y_var: str):
    """
    Run an OLS regression using statsmodels.

    Args:
      df: Input DataFrame
      x_vars: List of independent variable names
      y_var: Dependent variable name

    Returns:
      - statsmodels regression result object
    """
    X = df[x_vars]
    X = sm.add_constant(X)
    y = df[y_var]

    model = sm.OLS(y, X, missing="drop")
    res = model.fit()
    return res


def results_to_long(res, model_name: str) -> list[dict]:
    """
    Convert statsmodels regression results into
    a long-format list of dictionaries for CSV export.
    """
    rows = []
    for param in res.params.index:
        rows.append({
            "model": model_name,
            "variable": param,
            "coef": res.params[param],
            "std_err": res.bse[param],
            "t_stat": res.tvalues[param],
            "p_value": res.pvalues[param],
        })
    return rows


def apa_table_to_docx(results, labels, out_path: Path):
    """
    Create an APA-style regression table in a Word document.

    Args:
      results: list of statsmodels result objects
      labels: list of model labels
      out_path: output .docx file path
    """
    doc = Document()
    doc.add_heading("Regression Results", level=1)

    for res, label in zip(results, labels):
        doc.add_paragraph(label)
        table = doc.add_table(rows=len(res.params) + 1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = "Variable"
        hdr[1].text = "Coef."
        hdr[2].text = "Std. Err."
        hdr[3].text = "t"
        hdr[4].text = "p"

        for i, param in enumerate(res.params.index, start=1):
            row = table.rows[i].cells
            row[0].text = param
            row[1].text = f"{res.params[param]:.4f}"
            row[2].text = f"{res.bse[param]:.4f}"
            row[3].text = f"{res.tvalues[param]:.2f}"
            row[4].text = f"{res.pvalues[param]:.3f}"

    doc.save(out_path)

# MAIN PIPELINE
def main():
    """
    Orchestrates the full EPI–financial–MBS regression pipeline.
    """
    print("Loading EPI datasets...")
    epi_q, epi_topics, epi_aspects = load_epi_data()

    print("Loading financial data...")
    fin = load_financial_data()

    print("Building quarterly metrics...")
    qtr_metrics = build_quarterly_metrics(epi_q, fin)
    qtr_metrics.to_csv(OUT_MBS_QTR, index=False)
    print(f"Saved quarterly metrics: {OUT_MBS_QTR}")

    print("Saving merged modeling dataset...")
    qtr_metrics.to_csv(OUT_MERGED, index=False)
    print(f"Saved merged dataset: {OUT_MERGED}")

    # REGRESSION MODELS
    print("Running regression models...")

    # Baseline model
    X_BASE = ["EPI_lag1"]
    res_base = run_regression(qtr_metrics, X_BASE, DEP_VAR)

    # Extended model
    X_EXT = ["EPI_lag1", "EPI"]
    res_ext = run_regression(qtr_metrics, X_EXT, DEP_VAR)

    # EXPORT RESULTS
    print("Exporting regression results...")

    long = pd.DataFrame(
        results_to_long(res_base, "Baseline")
        + results_to_long(res_ext, "Extended")
    )
    long.to_csv(OUT_RESULTS, index=False)
    print(f"Saved regression results (long CSV): {OUT_RESULTS}")

    apa_table_to_docx(
        [res_base, res_ext],
        ["Model 1 (Baseline)", "Model 2 (Extended)"],
        OUT_TABLE_DOCX,
    )
    print(f"Saved APA-style regression table: {OUT_TABLE_DOCX}")

    # SANITY CHECK PRINTS
    print("\nRegression sample quarters used:")
    print(qtr_metrics[["quarter", "EPI_lag1", DEP_VAR]].head(10))

    print("\nN baseline:", int(res_base.nobs), " | N extended:", int(res_ext.nobs))
    print("Note: Q1 observations drop out because Q4 is not available to create a true q−1 lag.")
